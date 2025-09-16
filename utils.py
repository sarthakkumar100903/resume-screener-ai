# Enhanced utils.py — Multi-format Resume Parsing (PDF, DOCX, DOC), Embeddings, Contact Extraction, Azure Uploads

import re
import fitz  # PyMuPDF
import docx
import zipfile
from docx import Document
import numpy as np
import tiktoken
import functools
import asyncio
import logging
import time
from typing import List, Dict, Any, Optional, Tuple
from azure.storage.blob import BlobClient, BlobServiceClient
from sklearn.metrics.pairwise import cosine_similarity
from constants import AZURE_CONFIG, MODEL_CONFIG, PERFORMANCE_CONFIG
from openai import AzureOpenAI
import pandas as pd
import io

# Configure logging
logger = logging.getLogger(__name__)

# Initialize OpenAI client with optimized settings
client = AzureOpenAI(
    api_key=AZURE_CONFIG["openai_key"],
    api_version=AZURE_CONFIG["api_version"],
    azure_endpoint=AZURE_CONFIG["azure_endpoint"],
    max_retries=PERFORMANCE_CONFIG["max_retries"],
    timeout=PERFORMANCE_CONFIG["request_timeout"]
)

# ==========================
# 📄 Enhanced Multi-format Resume Text Extractor
# ==========================

def parse_resume(file_bytes: bytes, filename: str = "resume", max_pages: int = 10) -> str:
    """
    Enhanced resume parser supporting PDF, DOCX, and DOC formats
    """
    try:
        start_time = time.time()
        
        # Determine file type from filename or content
        file_extension = filename.lower().split('.')[-1] if '.' in filename else 'pdf'
        
        if file_extension == 'pdf':
            text = parse_pdf(file_bytes, max_pages)
        elif file_extension in ['docx', 'doc']:
            text = parse_word_document(file_bytes, file_extension)
        else:
            # Fallback: try to detect format by content
            text = parse_with_format_detection(file_bytes, max_pages)
        
        # Clean up text
        text = clean_resume_text(text)
        
        processing_time = time.time() - start_time
        logger.debug(f"Parsed {file_extension.upper()} in {processing_time:.2f}s, {len(text)} chars")
        
        return text
        
    except Exception as e:
        logger.error(f"Resume parsing failed for {filename}: {str(e)}")
        return f"Error reading resume: {str(e)}"

def parse_pdf(file_bytes: bytes, max_pages: int = 10) -> str:
    """Parse PDF files using PyMuPDF"""
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text_parts = []
            pages_processed = 0
            
            for page_num, page in enumerate(doc):
                if pages_processed >= max_pages:
                    break
                    
                try:
                    page_text = page.get_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                        pages_processed += 1
                except Exception as e:
                    logger.warning(f"Error processing PDF page {page_num}: {str(e)}")
                    continue
            
            return "\n\n".join(text_parts)
    except Exception as e:
        logger.error(f"PDF parsing error: {str(e)}")
        raise

def parse_word_document(file_bytes: bytes, file_extension: str) -> str:
    """Parse DOCX and DOC files"""
    try:
        if file_extension == 'docx':
            return parse_docx(file_bytes)
        elif file_extension == 'doc':
            return parse_doc_fallback(file_bytes)
        else:
            raise ValueError(f"Unsupported Word format: {file_extension}")
    except Exception as e:
        logger.error(f"Word document parsing error: {str(e)}")
        raise

def parse_docx(file_bytes: bytes) -> str:
    """Parse DOCX files using python-docx"""
    try:
        # Create a BytesIO object from bytes
        doc_io = io.BytesIO(file_bytes)
        doc = Document(doc_io)
        
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
        
    except Exception as e:
        logger.error(f"DOCX parsing error: {str(e)}")
        # Fallback: try to extract as ZIP
        return extract_docx_as_zip(file_bytes)

def extract_docx_as_zip(file_bytes: bytes) -> str:
    """Fallback method to extract DOCX content as ZIP"""
    try:
        doc_io = io.BytesIO(file_bytes)
        
        with zipfile.ZipFile(doc_io, 'r') as zip_file:
            # Try to read the main document content
            try:
                content = zip_file.read('word/document.xml')
                # Basic XML content extraction (remove tags)
                text = re.sub(r'<[^>]+>', ' ', content.decode('utf-8', errors='ignore'))
                text = re.sub(r'\s+', ' ', text).strip()
                return text
            except KeyError:
                # If document.xml is not found, try other files
                text_parts = []
                for filename in zip_file.namelist():
                    if filename.startswith('word/') and filename.endswith('.xml'):
                        try:
                            content = zip_file.read(filename)
                            clean_text = re.sub(r'<[^>]+>', ' ', content.decode('utf-8', errors='ignore'))
                            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
                            if clean_text:
                                text_parts.append(clean_text)
                        except Exception:
                            continue
                return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"DOCX ZIP extraction failed: {str(e)}")
        return "Error extracting DOCX content"

def parse_doc_fallback(file_bytes: bytes) -> str:
    """
    Fallback parser for DOC files (limited support)
    Note: Full DOC parsing requires additional libraries like python-docx2txt or antiword
    """
    try:
        # Basic text extraction attempt
        text = file_bytes.decode('utf-8', errors='ignore')
        
        # Remove common binary artifacts
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F-\xFF]', ' ', text)
        text = re.sub(r'\s+', ' ', text).strip()
        
        # If the text is mostly gibberish, return a message
        readable_chars = len(re.findall(r'[a-zA-Z\s]', text))
        if readable_chars < len(text) * 0.3:  # Less than 30% readable characters
            return "DOC format detected but content extraction is limited. Please convert to DOCX or PDF for better parsing."
        
        return text
        
    except Exception as e:
        logger.error(f"DOC parsing error: {str(e)}")
        return "Error parsing DOC file. Please convert to DOCX or PDF format."

def parse_with_format_detection(file_bytes: bytes, max_pages: int = 10) -> str:
    """Detect format and parse accordingly"""
    try:
        # Check for PDF signature
        if file_bytes.startswith(b'%PDF'):
            return parse_pdf(file_bytes, max_pages)
        
        # Check for DOCX signature (ZIP file with specific content)
        elif file_bytes.startswith(b'PK\x03\x04'):
            try:
                # Try to parse as DOCX first
                return parse_docx(file_bytes)
            except:
                # If DOCX parsing fails, might be another ZIP format
                return "Unsupported ZIP-based document format"
        
        # Check for DOC signature
        elif file_bytes.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            return parse_doc_fallback(file_bytes)
        
        else:
            # Try as text file
            try:
                return file_bytes.decode('utf-8', errors='ignore')
            except:
                return "Unknown file format - could not extract text"
                
    except Exception as e:
        logger.error(f"Format detection failed: {str(e)}")
        return f"Error detecting file format: {str(e)}"

def clean_resume_text(text: str) -> str:
    """Clean and normalize resume text"""
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # Max 2 consecutive newlines
    text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces and tabs
    
    # Remove common document artifacts
    text = re.sub(r'[^\x20-\x7E\n]', ' ', text)  # Remove non-ASCII except newlines
    text = re.sub(r'(\w)\s+(\w)', r'\1 \2', text)  # Fix broken words
    
    # Clean up common Word/PDF artifacts
    text = re.sub(r'\x0c', ' ', text)  # Form feed characters
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', ' ', text)  # Control characters
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    return text.strip()

# ==========================
# 📎 Enhanced Text Chunking
# ==========================

@functools.lru_cache(maxsize=1)
def get_tokenizer():
    """Cached tokenizer for performance"""
    return tiktoken.encoding_for_model("gpt-4")

def chunk_text(text: str, max_tokens: int = None, overlap: int = None) -> List[str]:
    """
    Enhanced text chunking with configurable parameters
    """
    if max_tokens is None:
        max_tokens = PERFORMANCE_CONFIG["chunk_size"]
    if overlap is None:
        overlap = PERFORMANCE_CONFIG["chunk_overlap"]
    
    if not text:
        return [""]
    
    try:
        enc = get_tokenizer()
        tokens = enc.encode(text)
        
        if len(tokens) <= max_tokens:
            return [text]
        
        chunks = []
        i = 0
        while i < len(tokens):
            chunk_tokens = tokens[i:i + max_tokens]
            chunk_text = enc.decode(chunk_tokens)
            chunks.append(chunk_text)
            i += max_tokens - overlap
            
            # Limit number of chunks for performance
            if len(chunks) >= PERFORMANCE_CONFIG["max_resume_chunks"]:
                break
        
        return chunks
        
    except Exception as e:
        logger.error(f"Text chunking failed: {str(e)}")
        return [text[:3000]]  # Fallback to simple truncation

def get_text_chunks(text: str, max_tokens: int = 800, overlap: int = 100) -> List[str]:
    """Wrapper for backward compatibility"""
    return chunk_text(text, max_tokens, overlap)

# ==========================
# 🧠 Enhanced Embedding & Similarity
# ==========================

@functools.lru_cache(maxsize=100)  # Increased cache size
def get_embedding_cached(text: str) -> Tuple[float, ...]:
    """
    Cached embedding generation with error handling and performance optimization
    """
    try:
        # Truncate text for faster embedding generation
        if len(text) > 8000:  # Embedding model limit is ~8191 tokens
            text = text[:8000]
        
        start_time = time.time()
        
        response = client.embeddings.create(
            input=[text],
            model=MODEL_CONFIG["embedding_model"]
        )
        
        embedding = response.data[0].embedding
        processing_time = time.time() - start_time
        
        logger.debug(f"Generated embedding in {processing_time:.2f}s")
        
        return tuple(embedding)  # Convert to tuple for caching
        
    except Exception as e:
        logger.error(f"Embedding generation failed: {str(e)}")
        # Return a zero vector as fallback
        return tuple([0.0] * 1536)

def get_embedding(text: str) -> List[float]:
    """Non-cached embedding for compatibility"""
    return list(get_embedding_cached(text))

def get_cosine_similarity(vec1: Tuple[float, ...], vec2: Tuple[float, ...]) -> float:
    """
    Enhanced cosine similarity calculation with validation
    """
    try:
        if not vec1 or not vec2:
            return 0.0
            
        if len(vec1) != len(vec2):
            logger.warning(f"Vector dimension mismatch: {len(vec1)} vs {len(vec2)}")
            return 0.0
        
        # Convert tuples to numpy arrays for calculation
        v1 = np.array(vec1).reshape(1, -1)
        v2 = np.array(vec2).reshape(1, -1)
        
        # Handle zero vectors
        if np.allclose(v1, 0) or np.allclose(v2, 0):
            return 0.0
        
        similarity = cosine_similarity(v1, v2)[0][0]
        
        # Ensure result is in valid range
        return max(0.0, min(1.0, float(similarity)))
        
    except Exception as e:
        logger.error(f"Cosine similarity calculation failed: {str(e)}")
        return 0.0

# ==========================
# 📬 Enhanced Contact Info Extractor
# ==========================

def extract_contact_info(text: str) -> Dict[str, str]:
    """
    Enhanced contact information extraction with multiple fallback strategies
    """
    if not text:
        return {"name": "N/A", "email": "N/A", "phone": "N/A"}
    
    contact_info = {
        "name": "N/A",
        "email": "N/A", 
        "phone": "N/A"
    }
    
    # Extract email with improved pattern
    email_patterns = [
        r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',  # Standard email
        r'[A-Za-z0-9._%+-]+\s*@\s*[A-Za-z0-9.-]+\s*\.\s*[A-Za-z]{2,}',  # Email with spaces
    ]
    
    for pattern in email_patterns:
        email_match = re.search(pattern, text)
        if email_match:
            contact_info["email"] = email_match.group(0).replace(" ", "")
            break
    
    # Extract phone with improved patterns
    phone_patterns = [
        r'\+?\d{1,3}[-.\s]?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',  # Various formats
        r'\+?\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}',        # International
        r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',                   # US format
        r'\+?\d{10,15}',                                          # Simple digits
    ]
    
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0)
            # Clean phone number
            phone = re.sub(r'[^\d+]', '', phone)
            if len(phone) >= 10:
                contact_info["phone"] = phone
                break
    
    # Enhanced name extraction
    contact_info["name"] = extract_candidate_name(text)
    
    return contact_info

def extract_candidate_name(text: str) -> str:
    """
    Enhanced name extraction with multiple strategies
    """
    if not text:
        return "N/A"
    
    lines = text.strip().split('\n')
    
    # Strategy 1: Look for name in first few lines
    for i, line in enumerate(lines[:5]):
        line = line.strip()
        if not line:
            continue
            
        # Skip lines with common resume headers
        skip_patterns = [
            r'resume|cv|curriculum|vitae|profile|objective|summary',
            r'contact|information|details|phone|email|address',
            r'experience|education|skills|projects|work|employment'
        ]
        
        if any(re.search(pattern, line, re.I) for pattern in skip_patterns):
            continue
        
        # Look for capitalized words that could be names
        words = line.split()
        if 2 <= len(words) <= 4:  # Typical name length
            if all(word[0].isupper() for word in words if word.isalpha()):
                # Additional validation
                if not any(char.isdigit() for char in line):  # No numbers in name
                    return line.strip()
    
    # Strategy 2: Look for "Name:" pattern
    name_match = re.search(r'name\s*:?\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)', text, re.I)
    if name_match:
        return name_match.group(1).strip()
    
    # Strategy 3: Find first line with proper capitalization
    for line in lines[:10]:
        line = line.strip()
        if len(line.split()) == 2:  # Likely first and last name
            words = line.split()
            if all(word[0].isupper() and word[1:].islower() for word in words):
                return line
    
    return "N/A"

# ==========================
# ☁️ Enhanced Azure Uploads with Error Handling
# ==========================

def upload_to_blob(file_bytes: bytes, file_name: str, container: str, 
                  overwrite: bool = True, max_retries: int = 3) -> bool:
    """
    Enhanced blob upload with retry logic and error handling
    """
    if not file_bytes or not file_name:
        logger.error("Invalid file data for blob upload")
        return False
    
    for attempt in range(max_retries):
        try:
            blob = BlobClient.from_connection_string(
                conn_str=AZURE_CONFIG["connection_string"],
                container_name=container,
                blob_name=file_name
            )
            
            blob.upload_blob(file_bytes, overwrite=overwrite)
            logger.debug(f"Successfully uploaded {file_name} to {container}")
            return True
            
        except Exception as e:
            logger.warning(f"Blob upload attempt {attempt + 1} failed: {str(e)}")
            if attempt < max_retries - 1:
                time.sleep(1.0 * (attempt + 1))  # Exponential backoff
            else:
                logger.error(f"Failed to upload {file_name} after {max_retries} attempts")
                return False
    
    return False

def save_summary_to_blob(pdf_bytes: bytes, file_name: str, container: str) -> bool:
    """Save PDF summary to blob storage"""
    return upload_to_blob(pdf_bytes, file_name, container)

def save_csv_to_blob(df: pd.DataFrame, file_name: str, container: str) -> bool:
    """Save CSV data to blob storage"""
    try:
        csv_data = df.to_csv(index=False).encode('utf-8')
        return upload_to_blob(csv_data, file_name, container)
    except Exception as e:
        logger.error(f"Failed to convert DataFrame to CSV: {str(e)}")
        return False

def download_from_blob(file_name: str, container: str) -> Optional[bytes]:
    """Download file from blob storage"""
    try:
        blob = BlobClient.from_connection_string(
            conn_str=AZURE_CONFIG["connection_string"],
            container_name=container,
            blob_name=file_name
        )
        
        return blob.download_blob().readall()
        
    except Exception as e:
        logger.error(f"Failed to download {file_name} from {container}: {str(e)}")
        return None

# ==========================
# 🔧 Utility Functions
# ==========================

def validate_resume_content(text: str) -> Dict[str, Any]:
    """
    Validate resume content and return quality metrics
    """
    if not text:
        return {
            "is_valid": False,
            "issues": ["Empty resume content"],
            "word_count": 0,
            "has_contact": False
        }
    
    word_count = len(text.split())
    issues = []
    
    # Check minimum length
    if len(text) < 100:
        issues.append("Resume too short")
    
    # Check maximum length  
    if len(text) > 50000:
        issues.append("Resume too long")
    
    # Check for contact information
    contact = extract_contact_info(text)
    has_contact = contact["email"] != "N/A" or contact["phone"] != "N/A"
    
    if not has_contact:
        issues.append("Missing contact information")
    
    # Check for suspicious patterns
    suspicious_patterns = [
        r'lorem ipsum', r'sample text', r'placeholder', 
        r'your name here', r'example\.com'
    ]
    
    for pattern in suspicious_patterns:
        if re.search(pattern, text, re.I):
            issues.append("Contains placeholder text")
            break
    
    return {
        "is_valid": len(issues) == 0,
        "issues": issues,
        "word_count": word_count,
        "character_count": len(text),
        "has_contact": has_contact,
        "contact_info": contact
    }

def get_file_format_from_name(filename: str) -> str:
    """Extract file format from filename"""
    if '.' in filename:
        return filename.lower().split('.')[-1]
    return 'unknown'

def is_supported_resume_format(filename: str) -> bool:
    """Check if file format is supported for resume parsing"""
    supported_formats = ['pdf', 'docx', 'doc']
    file_format = get_file_format_from_name(filename)
    return file_format in supported_formats

def format_processing_time(seconds: float) -> str:
    """Format processing time in human-readable format"""
    if seconds < 1:
        return f"{seconds*1000:.0f}ms"
    elif seconds < 60:
        return f"{seconds:.1f}s"
    else:
        minutes = int(seconds // 60)
        remaining_seconds = seconds % 60
        return f"{minutes}m {remaining_seconds:.1f}s"

def calculate_throughput(resume_count: int, total_time: float) -> float:
    """Calculate processing throughput in resumes per hour"""
    if total_time <= 0:
        return 0.0
    return (resume_count / total_time) * 3600

# ==========================
# 📊 Enhanced File Format Statistics
# ==========================

def analyze_file_formats(file_list: List[Tuple[str, bytes]]) -> Dict[str, Any]:
    """Analyze file formats in a list of files"""
    format_stats = {}
    total_size = 0
    
    for filename, file_bytes in file_list:
        file_format = get_file_format_from_name(filename)
        file_size = len(file_bytes)
        
        if file_format not in format_stats:
            format_stats[file_format] = {
                'count': 0,
                'total_size': 0,
                'files': []
            }
        
        format_stats[file_format]['count'] += 1
        format_stats[file_format]['total_size'] += file_size
        format_stats[file_format]['files'].append(filename)
        total_size += file_size
    
    # Calculate percentages
    for format_info in format_stats.values():
        format_info['size_percentage'] = (format_info['total_size'] / total_size) * 100 if total_size > 0 else 0
    
    return {
        'format_breakdown': format_stats,
        'total_files': len(file_list),
        'total_size_bytes': total_size,
        'total_size_mb': total_size / (1024 * 1024),
        'supported_formats': [fmt for fmt in format_stats.keys() if fmt in ['pdf', 'docx', 'doc']],
        'unsupported_formats': [fmt for fmt in format_stats.keys() if fmt not in ['pdf', 'docx', 'doc']]
    }

# ==========================
# 🔍 Advanced Text Processing for Multi-format Support
# ==========================

def extract_skills_from_text(text: str) -> List[str]:
    """Extract potential skills from resume text"""
    # Common technical skills patterns
    skill_patterns = [
        r'\b(?:Python|Java|JavaScript|C\+\+|C#|Ruby|PHP|Go|Rust|Kotlin|Swift)\b',
        r'\b(?:React|Angular|Vue|Node\.js|Django|Flask|Spring|Laravel)\b',
        r'\b(?:AWS|Azure|GCP|Docker|Kubernetes|Jenkins|Git|Linux|Windows)\b',
        r'\b(?:SQL|MySQL|PostgreSQL|MongoDB|Redis|Elasticsearch)\b',
        r'\b(?:HTML|CSS|REST|GraphQL|API|JSON|XML)\b'
    ]
    
    skills = set()
    for pattern in skill_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        skills.update(matches)
    
    return sorted(list(skills))

def extract_experience_years(text: str) -> Optional[int]:
    """Extract years of experience from resume text"""
    patterns = [
        r'(\d+)[\+\s]*years?\s+(?:of\s+)?experience',
        r'(\d+)[\+\s]*yrs?\s+(?:of\s+)?experience',
        r'experience[:\s]+(\d+)[\+\s]*years?',
        r'(\d+)[\+\s]*years?\s+in\s+(?:the\s+)?field'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            try:
                return int(match.group(1))
            except ValueError:
                continue
    
    return None

def extract_education_level(text: str) -> str:
    """Extract highest education level from resume text"""
    education_patterns = [
        (r'\b(?:PhD|Ph\.D|Doctorate|Doctoral)\b', 'PhD'),
        (r'\b(?:Masters?|M\.S|M\.A|MBA|M\.Tech|M\.E)\b', 'Masters'),
        (r'\b(?:Bachelor|B\.S|B\.A|B\.Tech|B\.E)\b', 'Bachelors'),
        (r'\b(?:Associate|A\.S|A\.A)\b', 'Associates'),
        (r'\b(?:High School|Diploma|12th)\b', 'High School')
    ]
    
    for pattern, level in education_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return level
    
    return "Not Specified"

# ==========================
# 📈 Quality Scoring with Multi-format Considerations
# ==========================

def calculate_resume_quality_score(text: str, contact: Dict[str, str], file_format: str = "pdf") -> Dict[str, Any]:
    """Calculate overall resume quality score with format-specific adjustments"""
    if not text:
        return {"score": 0, "factors": ["Empty resume"], "format": file_format}
    
    factors = []
    score = 0
    
    # Format-specific adjustments
    format_bonus = {
        'pdf': 0,    # Standard format
        'docx': 0,   # Standard format  
        'doc': -5    # Older format, slight penalty
    }
    
    score += format_bonus.get(file_format, -10)  # Unknown formats get penalty
    
    # Length check (10 points)
    word_count = len(text.split())
    if 200 <= word_count <= 2000:
        score += 10
    elif word_count < 200:
        factors.append("Too short")
    else:
        factors.append("Too long")
    
    # Contact information (20 points)
    contact_score = 0
    if contact.get("name", "N/A") != "N/A":
        contact_score += 7
    if contact.get("email", "N/A") != "N/A":
        contact_score += 7
    if contact.get("phone", "N/A") != "N/A":
        contact_score += 6
    score += contact_score
    
    if contact_score < 20:
        factors.append("Missing contact information")
    
    # Structure check (20 points)
    structure_indicators = [
        r'(?:experience|work|employment)',
        r'(?:education|academic)',
        r'(?:skills|competenc)',
        r'(?:project|portfolio)'
    ]
    
    structure_score = 0
    for indicator in structure_indicators:
        if re.search(indicator, text, re.IGNORECASE):
            structure_score += 5
    
    score += min(structure_score, 20)
    if structure_score < 15:
        factors.append("Poor structure")
    
    # Content quality (30 points)
    content_indicators = [
        (r'\b\d{4}\b', 5, "Has dates"),  # Years/dates
        (r'[%\d]+%|\d+\%', 5, "Has metrics"),  # Percentages/metrics
        (r'(?:achieved|improved|increased|developed|created|led|managed)', 10, "Action words"),
        (r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*\s+(?:Inc|Corp|LLC|Ltd|Company)', 10, "Company names")
    ]
    
    content_score = 0
    for pattern, points, description in content_indicators:
        if re.search(pattern, text, re.IGNORECASE):
            content_score += points
    
    score += min(content_score, 30)
    if content_score < 20:
        factors.append("Lacks specific details")
    
    # Technical content (20 points)
    skills = extract_skills_from_text(text)
    tech_score = min(len(skills) * 2, 20)
    score += tech_score
    
    if tech_score < 10:
        factors.append("Limited technical skills mentioned")
    
    return {
        "score": min(max(score, 0), 100),  # Ensure 0-100 range
        "word_count": word_count,
        "skills_found": skills,
        "factors": factors,
        "contact_completeness": contact_score / 20 * 100,
        "file_format": file_format,
        "format_quality": "Good" if file_format in ['pdf', 'docx'] else "Acceptable" if file_format == 'doc' else "Poor"
    }

# ==========================
# 🚨 Enhanced Fraud Detection with Multi-format Support
# ==========================

def detect_resume_anomalies(text: str, contact: Dict[str, str], file_format: str = "unknown") -> Dict[str, Any]:
    """Detect potential resume fraud or anomalies with format-specific checks"""
    anomalies = []
    risk_score = 0
    
    if not text:
        return {"anomalies": ["Empty resume"], "risk_score": 100, "format": file_format}
    
    # Format-specific risk assessment
    if file_format == 'doc':
        # DOC files have limited parsing, higher risk of incomplete analysis
        risk_score += 5
        anomalies.append("DOC format may have incomplete text extraction")
    elif file_format not in ['pdf', 'docx', 'doc']:
        risk_score += 15
        anomalies.append("Unusual file format for resumes")
    
    # Check for template indicators
    template_indicators = [
        r'lorem ipsum', r'sample text', r'placeholder',
        r'your name here', r'example\.com', r'\[.*\]',
        r'insert.*here', r'replace.*with'
    ]
    
    for pattern in template_indicators:
        if re.search(pattern, text, re.IGNORECASE):
            anomalies.append("Contains template/placeholder text")
            risk_score += 20
            break
    
    # Check for excessive claims
    excessive_claims = [
        r'expert in (?:all|every|100\+)',
        r'single-handedly (?:increased|improved|developed)',
        r'(?:increased|improved).*(?:1000|500)%',
        r'worked with (?:all|every) Fortune',
        r'(?:invented|pioneered|revolutionized)'
    ]
    
    for pattern in excessive_claims:
        if re.search(pattern, text, re.IGNORECASE):
            anomalies.append("Contains unrealistic claims")
            risk_score += 15
            break
    
    # Check for inconsistent information
    years_mentioned = re.findall(r'\b(19|20)\d{2}\b', text)
    if len(years_mentioned) > 10:  # Too many years mentioned
        anomalies.append("Excessive date references")
        risk_score += 10
    
    # Check contact information consistency
    emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
    if len(emails) > 3:
        anomalies.append("Multiple email addresses")
        risk_score += 10
    
    # Check for duplicate content
    sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 20]
    if len(sentences) != len(set(sentences)):
        anomalies.append("Contains duplicate content")
        risk_score += 15
    
    # Format-specific anomaly detection
    if file_format == 'pdf' and 'Error reading resume' in text:
        anomalies.append("PDF parsing errors detected")
        risk_score += 10
    elif file_format in ['docx', 'doc'] and len(text) < 100:
        anomalies.append("Word document appears to have minimal content")
        risk_score += 10
    
    return {
        "anomalies": anomalies,
        "risk_score": min(risk_score, 100),
        "requires_manual_review": risk_score > 30,
        "file_format": file_format,
        "format_specific_issues": [a for a in anomalies if 'format' in a.lower() or 'parsing' in a.lower()]
    }

# ==========================
# 📊 Export Utilities with Format Information
# ==========================

def prepare_export_data(df: pd.DataFrame, include_sensitive: bool = False) -> pd.DataFrame:
    """Prepare dataframe for export with optional sensitive data filtering"""
    export_df = df.copy()
    
    # Remove large text fields to reduce file size
    columns_to_remove = ["resume_text", "embedding"]
    if not include_sensitive:
        columns_to_remove.extend(["phone"])  # Remove phone for privacy
    
    export_df = export_df.drop(columns=columns_to_remove, errors='ignore')
    
    # Clean up list fields for CSV export
    list_columns = ["red_flags", "missing_gaps", "highlights", "reasons_if_rejected"]
    for col in list_columns:
        if col in export_df.columns:
            export_df[col] = export_df[col].apply(
                lambda x: "; ".join(x) if isinstance(x, list) else str(x)
            )
    
    # Add file format information if available
    if 'resume_file' in export_df.columns:
        export_df['file_format'] = export_df['resume_file'].apply(get_file_format_from_name)
    
    return export_df

def generate_analysis_summary(results: List[Dict[str, Any]]) -> Dict[str, Any]:
    """Generate comprehensive analysis summary with format breakdown"""
    if not results:
        return {"error": "No results to analyze"}
    
    df = pd.DataFrame(results)
    
    # File format analysis
    format_breakdown = {}
    if 'resume_file' in df.columns:
        df['file_format'] = df['resume_file'].apply(get_file_format_from_name)
        format_breakdown = df['file_format'].value_counts().to_dict()
    
    summary = {
        "total_candidates": len(df),
        "average_score": df["score"].mean(),
        "score_distribution": {
            "excellent": len(df[df["score"] >= 80]),
            "good": len(df[(df["score"] >= 60) & (df["score"] < 80)]),
            "average": len(df[(df["score"] >= 40) & (df["score"] < 60)]),
            "poor": len(df[df["score"] < 40])
        },
        "verdict_distribution": df["verdict"].value_counts().to_dict(),
        "format_breakdown": format_breakdown,
        "top_skills": extract_top_skills_from_results(results),
        "quality_metrics": {
            "fraud_detected": len(df[df["fraud_detected"] == True]),
            "missing_contact": len(df[df["email"] == "N/A"]),
            "complete_profiles": len(df[(df["email"] != "N/A") & (df["phone"] != "N/A")])
        }
    }
    
    return summary

def extract_top_skills_from_results(results: List[Dict[str, Any]]) -> List[Dict[str, int]]:
    """Extract most mentioned skills across all resumes"""
    skill_counts = {}
    
    for result in results:
        resume_text = result.get("resume_text", "")
        if resume_text:
            skills = extract_skills_from_text(resume_text)
            for skill in skills:
                skill_counts[skill] = skill_counts.get(skill, 0) + 1
    
    # Sort by frequency and return top 10
    sorted_skills = sorted(skill_counts.items(), key=lambda x: x[1], reverse=True)
    return [{"skill": skill, "count": count} for skill, count in sorted_skills[:10]]
